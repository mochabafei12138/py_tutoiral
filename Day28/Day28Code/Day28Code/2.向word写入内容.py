# 注意1：我们在安装此模块使用的是pip install python-docx，但是在导入的时候是 docx
# 注意2：一般情况下，如果安装了A库，同时自动安装了依赖B库，如果A库更新了，如果使用过程中报错，则可以更新依赖库

# C:\Users\19621>pip install python-docx
# Collecting python-docx
#   Downloading python_docx-1.1.0-py3-none-any.whl (239 kB)
#      |████████████████████████████████| 239 kB 930 kB/s
# Requirement already satisfied: lxml>=3.1.0 in d:\software\anaconda\lib\site-packages (from python-docx) (4.6.3)
# Requirement already satisfied: typing-extensions in d:\software\anaconda\lib\site-packages (from python-docx) (4.9.0)
# Installing collected packages: python-docx
# Successfully installed python-docx-1.1.0

# 比如：安装python-docx的过程中，自动安装了依赖库lxml和typing-extensions

from docx import  Document
from docx.shared import Cm

# 1.创建doc文档对象
doc = Document()
# 2.添加标题
doc.add_heading('今天星期二~~~~')
# 3.添加段落
p = doc.add_paragraph('第一个段落~~~~')
# 注意：一个段落可以由很多个run组成
run1 = p.add_run('this is a test')
run2 = p.add_run('Python是一门面向对象的语言')
run3 = p.add_run('Python是跨平台的')

# 给每个run可以单独设置样式
run1.bold = True
run2.underline = True

# 4.添加图片
doc.add_picture(r'data/3.png',width=Cm(10))

# 5.添加列表
# a.有序列表
doc.add_paragraph('教学部',style='List Number')
doc.add_paragraph('运营部',style='List Number')
doc.add_paragraph('财务部',style='List Number')
doc.add_paragraph('行政部',style='List Number')
doc.add_paragraph('市场部',style='List Number')

# b.无序列表
doc.add_paragraph('教学部',style='List Bullet')
doc.add_paragraph('运营部',style='List Bullet')
doc.add_paragraph('财务部',style='List Bullet')
doc.add_paragraph('行政部',style='List Bullet')
doc.add_paragraph('市场部',style='List Bullet')

# 保存文件
doc.save(r'data/Python操作word.docx')


