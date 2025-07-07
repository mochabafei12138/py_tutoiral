from docx import  Document

# 1.打开doc
doc = Document(r'data/占勇辉的离职证明.docx')

# 2.获取doc中的所有内容，返回一个列表，其中的元素是段落对象
print(doc.paragraphs)

# 3.遍历列表，获取每个段落对象
for para in doc.paragraphs:
    # 获取每个段落的文本
    # print(para.text)
    # 获取组成每个段落的run，返回一个列表
    print(para.runs)
    for run in para.runs:
        # 获取每个run的文本
        print(run.text)
